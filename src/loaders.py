import os
import fitz  # pymupdf
from docx import Document as DocxDocument
from langchain_core.documents import Document
from src.config import DOCS_FOLDER, MAX_FILES

def load_txt(path: str) -> str:
    with open(path, "r") as f:
        return f.read()

def load_pdf(path: str) -> str:
    doc = fitz.open(path)
    return "\n".join([page.get_text() for page in doc])

def load_docx(path: str) -> str:
    doc = DocxDocument(path)
    return "\n".join([para.text for para in doc.paragraphs])

def load_file(path: str) -> str:
    ext = path.split(".")[-1].lower()
    loaders = {
        "txt": load_txt,
        "pdf": load_pdf,
        "docx": load_docx
    }
    if ext not in loaders:
        print(f"⚠️  Skipping unsupported file: {path}")
        return ""
    return loaders[ext](path)

def load_all_documents() -> list[Document]:
    files = [
        f for f in os.listdir(DOCS_FOLDER)
        if f.endswith((".txt", ".pdf", ".docx"))
    ][:MAX_FILES]

    if not files:
        print("❌ No documents found in docs/ folder!")
        exit()

    print(f"\n📂 Loading {len(files)} file(s):")
    all_docs = []

    for file in files:
        path = os.path.join(DOCS_FOLDER, file)
        print(f"   → {file}")
        text = load_file(path)
        if text:
            all_docs.append(Document(
                page_content=text,
                metadata={"source": file}
            ))

    return all_docs
